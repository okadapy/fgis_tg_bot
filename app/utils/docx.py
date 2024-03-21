import docx
from os import makedirs
from os.path import join
import datetime
import re
import qrcode
from docx.shared import Inches
from docx.text.paragraph import Paragraph

basic_re = r"\{[^\s}￼]*\}"
makedirs(join("media", "output_files"), exist_ok=True)


async def create_document(data: dict, vri_id: str, add_stamp: bool, add_header) -> str:
    doc = docx.Document(join("media", "doc-template.docx"))
    try:
        filename = join(
            "media",
            "output_files",
            data["vriInfo"]["applicable"]["certNum"].replace("/", "_") + ".docx",
        )
        qrcode_path = join(
            "media",
            "output_files",
            data["vriInfo"]["applicable"]["certNum"].replace("/", "_") + ".png",
        )
        stamp_path = join("media", "stamp.jpg")
        img = qrcode.make("https://fgis.gost.ru/fundmetrology/cm/results/" + vri_id)
        img.save(qrcode_path)

    except KeyError as ke:
        filename = join(
            "media",
            "output_files",
            data["vriInfo"]["inapplicable"]["certNum"].replace("/", "_") + ".docx",
        )
        qrcode_path = join(
            "media",
            "output_files",
            data["vriInfo"]["inapplicable"]["certNum"].replace("/", "_") + ".png",
        )
        stamp_path = join("media", "stamp.jpg")
        img = qrcode.make("https://fgis.gost.ru/fundmetrology/cm/results/" + vri_id)
        img.save(qrcode_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraph: Paragraph
                for paragraph in cell.paragraphs:
                    matches = re.findall(basic_re, paragraph.text)
                    if matches:
                        for match in matches:
                            if match == "{means/mis}":
                                line = []
                                if "mis" in data["means"].keys():
                                    for i in data["means"]["mis"]:
                                        line.append(
                                            f'{i["mitypeTitle"]} {i["number"]};'
                                        )
                                    paragraph.text = paragraph.text.replace(
                                        "{means/mis}", " ".join(line)
                                    )
                                else:
                                    paragraph.text = ""
                            elif match == "{means/mieta}":
                                line = []
                                for i in data["means"]["mieta"]:
                                    line.append(
                                        f'{i["regNumber"]} {i["mitypeTitle"]} {i["notation"]};'
                                    )
                                paragraph.text = paragraph.text.replace(
                                    "{means/mieta}", " ".join(line)
                                )
                            elif match == "{vriInfo/applicable}":
                                if any(data["vriInfo"]["applicable"]):
                                    paragraph.text = paragraph.text.replace(
                                        "{vriInfo/applicable}", "пригодным"
                                    )
                                else:
                                    paragraph.text = paragraph.text.replace(
                                        "{vriInfo/applicable}", "непригодным"
                                    )
                            elif match == "{vriInfo/vriType}":
                                if data["vriInfo"]["vriType"] == 1:
                                    paragraph.text = paragraph.text.replace(
                                        "{vriInfo/vriType}", "первичной"
                                    )
                                else:
                                    paragraph.text = paragraph.text.replace(
                                        "{vriInfo/vriType}", "периодической"
                                    )
                            elif match == "{vriId-no-prefix}":
                                paragraph.text = paragraph.text.replace(
                                    "{vriId-no-prefix}", vri_id[2:]
                                )
                            elif match == "{vriId}":
                                paragraph.text = paragraph.text.replace(
                                    "{vriId}", vri_id
                                )
                            elif match == "{qrCode}":
                                paragraph.text = paragraph.text.replace("{qrCode}", "")
                                run = paragraph.add_run()
                                run.add_picture(qrcode_path, width=Inches(0.5))
                            elif match == "{stamp}" and "2024" in data["vriInfo"]["vrfDate"]:
                                paragraph.text = paragraph.text.replace("{stamp}", "")
                                if add_stamp == True:
                                    run = paragraph.add_run()
                                    run.add_picture(stamp_path, width=Inches(0.5))
                            elif match == "{header_line_1}":
                                if add_header:
                                    paragraph.text = paragraph.text.replace(
                                        "{header_line_1}",
                                        'Общество с ограниченной ответственностью "КВАЗАР"',
                                    )
                                else:
                                    paragraph.text = ""
                            elif match == "{header_line_2}":
                                if add_header:
                                    paragraph.text = paragraph.text.replace(
                                        "{header_line_2}",
                                        "108823, Москва г, Рязановское п, Знамя Октября п, дом 31,  этаж 1,  пом. 38",
                                    )
                                else:
                                    paragraph.text = ""
                            elif match == "{header_line_3}":
                                if add_header:
                                    paragraph.text = paragraph.text.replace(
                                        "{header_line_3}",
                                        "Уникальный номер записи об аккредитации в реестре аккредитованных лиц RA.RU.310696",
                                    )
                                else:
                                    paragraph.text = ""
                            else:
                                find_tree = (
                                    match.replace("{", "").replace("}", "").split("/")
                                )
                                value = data
                                for i in find_tree:
                                    value = value[i]
                                paragraph.text = paragraph.text.replace(match, value)

    doc.save(filename)
    return filename
